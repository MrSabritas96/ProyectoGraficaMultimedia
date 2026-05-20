import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def create_docx(md_path, docx_path):
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Documento de Proyecto: MedTrack', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if it's a table
        if line.startswith('|'):
            in_table = True
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if set(cells[0]) == {'-'} or '---' in cells[0]: # Separator
                continue
            table_data.append(cells)
            continue
        else:
            if in_table:
                # We finished collecting table data, let's render it
                render_table(doc, table_data)
                in_table = False
                table_data = []

        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('• '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            
    if in_table:
        render_table(doc, table_data)

    doc.save(docx_path)

def add_formatted_text(paragraph, text):
    # Handle bold and italics using basic regex splitting
    # Note: This is a simple parser and doesn't handle nested formatting perfectly
    
    # Find all bold and italic parts
    # First replace <br> with newlines if any
    text = text.replace('<br>', '\n')
    
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def render_table(doc, data):
    if not data:
        return
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for r in range(rows):
        row_cells = table.rows[r].cells
        for c in range(cols):
            # Clean up the <br> for table cells
            cell_text = data[r][c].replace('<br>', '\n')
            # If it's the header, make it bold
            p = row_cells[c].paragraphs[0]
            if r == 0:
                # Remove bold asterisks from header if they exist
                clean_text = cell_text.replace('**', '')
                run = p.add_run(clean_text)
                run.bold = True
            else:
                add_formatted_text(p, cell_text)

if __name__ == '__main__':
    create_docx('exel/documento_proyecto.md', 'exel/Proyecto_MedTrack_Oficial.docx')
    print("Convertido con exito.")
